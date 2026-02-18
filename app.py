import streamlit as st
import requests
from bs4 import BeautifulSoup
import io
import zipfile
import re
import docx
import pypdf
from docx import Document
from docx.oxml.shared import qn
from docx.oxml import OxmlElement
from docx.shared import Inches
import streamlit.components.v1 as components
import requests
import json

# Configuration de la page (DOIT être la première commande Streamlit)
st.set_page_config(page_title="Architecte des Transitions", page_icon="🏗️", layout="wide")
if 'fusion_generee' not in st.session_state:
    st.session_state.fusion_generee = False

def extract_text_from_file(uploaded_file):
    """Extrait le texte brut d'un fichier PDF, DOCX ou TXT."""
    text = ""
    try:
        if uploaded_file.name.endswith('.pdf'):
            reader = pypdf.PdfReader(uploaded_file)
            for page in reader.pages:
                text += page.extract_text() + "\n"
        elif uploaded_file.name.endswith('.docx'):
            doc = docx.Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
        elif uploaded_file.name.endswith('.txt'):
            text = uploaded_file.read().decode('utf-8')
    except Exception as e:
        return f"Erreur de lecture du fichier : {e}"
    return text

def add_hyperlink(paragraph, url, text):
    """Insère un hyperlien cliquable dans un paragraphe Word."""
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    hyperlink = OxmlElement('w:hyperlink')
    hyperlink.set(qn('r:id'), r_id)

    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Appliquer le style bleu et souligné
    c = OxmlElement('w:color')
    c.set(qn('w:val'), '0000FF')
    rPr.append(c)
    u = OxmlElement('w:u')
    u.set(qn('w:val'), 'single')
    rPr.append(u)

    new_run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)
    return hyperlink

def convert_html_to_doc_format(html_content):
    """Encapsule le contenu pour Word avec gestion des tableaux et style HTML conservé."""
    soup = BeautifulSoup(html_content, 'html.parser')
    html_header = (
        '<html xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w="urn:schemas-microsoft-com:office:word" '
        'xmlns="http://www.w3.org/TR/REC-html40"><head><meta charset="utf-8"></head><body>'
    )
    full_html = html_header + str(soup) + '</body></html>'
    return io.BytesIO(full_html.encode('utf-8'))

def build_constraints_block(data: dict) -> str:
    """Construit un bloc HTML structuré pour les contraintes (onglet 3)."""
    # Tableau simple en HTML pour garder une forme propre dans Word
    rows = []
    labels = {
        "Périmètre": "Périmètre réduit du problème",
        "Nature": "Type de problème",
        "Douleurs": "Douleurs perçues par les acteurs",
        "Partenaires": "Partenaires obligatoires / Compétiteurs",
        "Bénéfices T": "Bénéfices tangibles",
        "Bénéfices I": "Bénéfices intangibles",
        "Objectifs": "Objectifs chiffrés et opposables",
        "Budget": "Budget prévisionnel",
        "Planning": "Planning et jalons",
        "Com": "Communication / Visibilité",
        "Marketing": "Vecteurs marketing",
        "Infos": "Informations complémentaires"
    }
    for k, v in data.items():
        label = labels.get(k, k)
        content = v if v else "Non précisé"
        rows.append(
            f"<tr><th style='text-align:left;padding:4px 8px;'>{label}</th>"
            f"<td style='padding:4px 8px;'>{content}</td></tr>"
        )
    table_html = (
        "<h2>Cadrage du problème et contraintes</h2>"
        "<table border='1' style='border-collapse:collapse;width:100%;'>"
        f"{''.join(rows)}"
        "</table>"
    )
    return table_html

def merge_html_and_constraints(sofia_html: str, constraints_data: dict) -> io.BytesIO:
    """
    Fusionne le HTML de SofIA et le bloc de contraintes (en HTML)
    dans un seul fichier .doc utilisable par Word.
    """
    constraints_html = build_constraints_block(constraints_data)
    sofia_part = BeautifulSoup(sofia_html, 'html.parser')

    # On encapsule les deux sections dans un même body pour que Word garde la mise en forme HTML
    merged_html = (
        '<html xmlns:o="urn:schemas-microsoft-com:office:office" '
        'xmlns:w="urn:schemas-microsoft-com:office:word" '
        'xmlns="http://www.w3.org/TR/REC-html40">'
        '<head><meta charset="utf-8"></head><body>'
        '<h1>Analyse SofIA</h1>'
        f'{str(sofia_part)}'
        '<hr/>'
        f'{constraints_html}'
        '</body></html>'
    )
    return io.BytesIO(merged_html.encode('utf-8'))

# --- INITIALISATION SESSION STATE ---
if "messages" not in st.session_state:
    st.session_state.messages = []
if "cadrage" not in st.session_state:
    st.session_state.cadrage = {}
if 'prompt_genere' not in st.session_state:
    st.session_state.prompt_genere = False
if 'sofia_html' not in st.session_state:
    st.session_state.sofia_html = None

# --- HEADER (LOGO + TITRE PROJET) ---
col_logo, col_titre = st.columns([1, 5])

with col_logo:
    try:
        st.image("LOGO ARC.jpg", use_container_width=True)
    except:
        st.warning("Logo non trouvé")

with col_titre:
    st.markdown("""
        <div style='margin-top: 20px;'>
            <h2 style='color: #2E4053;'>Projet Exploratoire Formation IMT - l'Architecte des Transitions</h2>
        </div>
    """, unsafe_allow_html=True)

st.markdown("---")

# --- INTERFACE STREAMLIT ---
st.title("Assistant pour formuler un problématique et utiliser plusieurs Agents")

tab0, tab1, tab2, tab3, tab4, tab5, tab6 = st.tabs([
    "📝 0.Présentation du projet", 
    "📝 1.Aide au Prompt pour SofIA",
    "📂 2.Extraction & Export",
    "📝 3.Aide au Prompt Contraintes",
    "🤖 4.Eval IA",
    "🤖 5.Chaos IA",
    "📝 6.Questionnaire"
])

# =============================================================================
# SOLUTION LLM - Reformulation Intelligente du Prompt (onglet 1)
# =============================================================================
def build_prompt_with_llm(q1, q2, q3, q4, q5):
    """
    Utilise l'API Claude Haiku pour reformuler intelligemment le prompt.
    Gère automatiquement les champs vides et produit une phrase fluide.
    
    Si l'API échoue, retourne None (à gérer par l'appelant).
    """
    # Nettoyer et préparer les informations disponibles
    infos = {}
    if q1 and q1.strip():
        infos['objectif_principal'] = q1.strip()
    if q2 and q2.strip():
        infos['perimetre'] = q2.strip()
    if q3 and q3.strip():
        infos['cibles'] = q3.strip()
    if q4 and q4.strip():
        infos['objectif_chiffre'] = q4.strip()
    if q5 and q5.strip():
        infos['action_complementaire'] = q5.strip()
    
    # Si aucune info fournie
    if not infos:
        return None
    
    # Construire le prompt système pour Claude
    system_prompt = """Tu es un assistant expert qui aide à formuler des problématiques pour l'outil SofIA de l'ADEME.

Ton rôle est de créer une question de recherche cohérente et bien formulée à partir des informations fournies.

RÈGLES STRICTES :
1. Intègre TOUTES les informations disponibles de manière naturelle et fluide
2. La question doit être grammaticalement parfaite (élisions correctes, pas de redondances)
3. La formulation doit être professionnelle et adaptée à un contexte institutionnel
4. OBLIGATOIRE : Termine la question par cette série de questions systémiques (à copier exactement) :

"Quelles sont les principales données dans ce domaine, quelles sont les données dont disposent l'ADEME dans ce domaine, quels sont les acteurs à mobiliser, les paramètres clés à travailler. Quelles sont les solutions déjà mises en œuvre, les principaux résultats déjà obtenus, les projets ayant réussi, leurs résultats et ceux ayant échoué et leurs causes. Quelles sont les règles de fonctionnement du système considéré, les paradigmes du système considéré et comment le transcender pour réduire le problème et identifier de nouvelles solutions. Quels sont les effets et conséquences systémiques liés à ce problème et aux futures actions dans d'autres domaines, les recommandations pour intégrer les effets rebonds, boucles de rétroactions et cobénéfices ?"

IMPORTANT : Réponds UNIQUEMENT avec la question reformulée finale. Pas d'introduction, pas d'explication, juste la question."""

    user_prompt = f"""Informations disponibles :
{json.dumps(infos, indent=2, ensure_ascii=False)}

Reformule ces informations en une question cohérente pour SofIA."""

    try:
        response = requests.post(
            "https://api.anthropic.com/v1/messages",
            headers={
                "Content-Type": "application/json",
                "anthropic-version": "2023-06-01"
            },
            json={
                "model": "claude-haiku-4-20250514",
                "max_tokens": 1500,
                "system": system_prompt,
                "messages": [
                    {"role": "user", "content": user_prompt}
                ]
            },
            timeout=30
        )
        
        if response.status_code == 200:
            result = response.json()
            return result['content'][0]['text'].strip()
        else:
            print(f"Erreur API {response.status_code}: {response.text}")
            return None
            
    except Exception as e:
        print(f"Erreur lors de l'appel API : {e}")
        return None

# --- ONGLET 0 : PRésentation du projet ART ---
with tab0: 
    st.header("0.Présentation du projet ART")
    st.info("L'ARchitecte des Transitions, ART, est un projet exploratoire visant à augmenter les capacités d'actions de l'expert et du manager; L'objectif est de les aider dans la conception de leur mode d'intervention dans une forte incertitude technique, financière, sociale, donc systémique.")
    st.info("Pour cela, plusieurs Agents IA ont été développés sur la base de +100 évaluations de programmes réalisées ces 20 dernières années. Fort de ces connaissances, les Agents IA peuvent vous accompagner pour concevoir de nouveaux modes d'intervention, mieux intégrer les aspects systémiques, pour intégrer une série de contraintes et évaluer leur conséquence.")
    st.error("Pour toute question : Gabriel Plassat")
    st.image("incertitude_syst.png", caption="Schéma", use_container_width=True)
    
# --- ONGLET 1 : AIDE AU PROMPT "SofIA" ---
with tab1: 
    st.header("1.Aide pour formuler le problème initial à SofIA")
    st.info("SofIA va être utilisé pour rédiger la problématique complète. L'IA reformulera automatiquement vos réponses en une question cohérente, même si certains champs sont vides.")
    
    q1 = st.text_area(
        "1. Votre objectif principal :", 
        placeholder="ex: développer la pratique de la marche au quotidien",
        help="Décrivez votre objectif principal. Ce champ n'est pas obligatoire."
    )
    q2 = st.text_input(
        "2. Périmètre géographique :", 
        placeholder="ex: dans tous les territoires",
        help="Précisez le périmètre (optionnel)."
    )
    q3 = st.text_area(
        "3. Cibles visées en priorité :", 
        placeholder="ex: toutes les personnes à tous les âges",
        help="Identifiez les cibles prioritaires (optionnel)."
    )
    q4 = st.text_input(
        "4. Objectif chiffré :", 
        placeholder="ex: augmenter de 20% la part de la marche",
        help="Donnez un objectif mesurable si possible (optionnel)."
    )
    q5 = st.text_area(
        "5. Action complémentaire ?", 
        placeholder="ex: étudier plus particulièrement les trajets domicile-travail",
        help="Ajoutez des compléments d'information (optionnel)."
    )
    
    st.info("💡 Vous n'êtes pas obligé de remplir tous les champs ! L'IA s'adaptera automatiquement.")
    
    if st.button("🤖 Générer le prompt avec IA", type="primary"):
        # Vérifier qu'au moins un champ est rempli
        if not any([q1.strip(), q2.strip(), q3.strip(), q4.strip(), q5.strip()]):
            st.error("⚠️ Veuillez remplir au moins un champ avant de générer le prompt.")
        else:
            # Générer le prompt avec LLM
            with st.spinner("🤖 Reformulation intelligente en cours... Cela peut prendre quelques secondes."):
                phrase_prompt = build_prompt_with_llm(q1, q2, q3, q4, q5)
            
            if phrase_prompt is None:
                st.error("❌ Erreur lors de la génération du prompt. Vérifiez votre connexion et réessayez.")
            else:
                # Afficher un aperçu
                st.success("✅ Prompt généré avec succès !")
                
                with st.expander("👁️ Aperçu du prompt généré", expanded=True):
                    st.markdown(phrase_prompt)
                    st.caption("Vous pouvez relire et vérifier le prompt avant de le télécharger.")
                
                # Créer le document Word
                prompt_doc = Document()
                prompt_doc.add_heading("Prompt pour SofIA", 0)
           
                try:
                    prompt_doc.add_paragraph("Utilisez le prompt ci-dessous dans l'interface SofIA :")
                    prompt_doc.add_picture("sofia_q.png", width=Inches(5.5))
                except Exception as e:
                    # Image non trouvée, on continue sans
                    pass
                
                prompt_doc.add_heading("Votre prompt personnalisé (généré par IA) :", level=1)
                prompt_doc.add_paragraph(phrase_prompt)
                
                prompt_doc.add_heading("Lien vers SofIA : https://www.sofia-transition-ecologique.fr/", level=1)
                p = prompt_doc.add_paragraph("Copiez/collez ce prompt dans SofIA. Se connecter à ")
                add_hyperlink(p, "https://www.sofia-transition-ecologique.fr/", "SofIA")
                
                # Sauvegarder
                prompt_buffer = io.BytesIO()
                prompt_doc.save(prompt_buffer)
                prompt_buffer.seek(0)
                
                st.download_button(
                    label="📥 Télécharger votre prompt pour SofIA",
                    data=prompt_buffer,
                    file_name="Prompt_Initial_Sofia.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    on_click=lambda: st.session_state.update({"prompt_genere": True})
                )

                if st.session_state.get('prompt_genere', False):
                    st.markdown("---")
                    st.success("✅ Document téléchargé !")
                    st.markdown("### 🚀 Étape suivante : Connectez-vous à [SofIA](https://www.sofia-transition-ecologique.fr/)")

    # Aide contextuelle
    with st.expander("ℹ️ Comment remplir les champs ?"):
        st.markdown("""
        **Guide rapide :**
        
        - **Objectif principal** : Ce que vous voulez accomplir (ex: "développer la mobilité douce")
        - **Périmètre** : Où (ex: "en Île-de-France" ou "sur tout le territoire national")
        - **Cibles** : Pour qui (ex: "les jeunes de 18-25 ans" ou "les entreprises de plus de 50 salariés")
        - **Objectif chiffré** : Un objectif mesurable (ex: "réduire de 30% les émissions")
        - **Action complémentaire** : Détails supplémentaires (ex: "tout en préservant l'emploi local")
        
        💡 **Important :** Vous n'avez pas besoin de remplir tous les champs ! L'IA reformulera intelligemment 
        les informations que vous fournissez, même partielles.
        
        🤖 **Avantage de l'IA :** La reformulation sera grammaticalement correcte, fluide et professionnelle, 
        quelle que soit la façon dont vous remplissez les champs.
        """)

    st.image("sofia_q.png", caption="Interface SofIA", use_container_width=True)
"""

# =============================================================================
# TESTS
# =============================================================================

if __name__ == "__main__":
    print("=== TEST DE LA SOLUTION LLM ===\n")
    
    # Test 1 : Tous les champs
    print("Test 1 : Tous les champs remplis")
    print("-" * 80)
    result1 = build_prompt_with_llm(
        "développer la pratique de la marche au quotidien",
        "dans tous les territoires",
        "toutes les personnes à tous les âges",
        "augmenter de 20% la part de la marche",
        "étudier plus particulièrement les trajets domicile-travail"
    )
    if result1:
        print(result1)
    else:
        print("❌ Erreur lors de la génération")
    print("\n" + "=" * 80 + "\n")
    
    # Test 2 : Seulement 2 champs
    print("Test 2 : Seulement objectif et cibles")
    print("-" * 80)
    result2 = build_prompt_with_llm(
        "réduire la consommation d'énergie des bâtiments",
        "",
        "les copropriétés de plus de 50 lots",
        "",
        ""
    )
    if result2:
        print(result2)
    else:
        print("❌ Erreur lors de la génération")
    print("\n" + "=" * 80 + "\n")
    
    # Test 3 : Un seul champ
    print("Test 3 : Seulement objectif chiffré")
    print("-" * 80)
    result3 = build_prompt_with_llm(
        "",
        "",
        "",
        "atteindre 100 000 véhicules électriques en circulation",
        ""
    )
    if result3:
        print(result3)
    else:
        print("❌ Erreur lors de la génération")
    print("\n" + "=" * 80 + "\n")


# --- ONGLET 2 : EXTRACTION ---
with tab2:
    st.header("2.Récupération de l'exportation")
    st.info("Importez le fichier chat_history.html généré par SofIA pour le transformer en fichier au format docx")
    uploaded_file = st.file_uploader("Glissez votre fichier chat_history.html ici", type="html", key="uploader")

    if uploaded_file:
        content = uploaded_file.read().decode('utf-8')
        # On stocke le HTML pour réutilisation dans l'onglet 4
        st.session_state.sofia_html = content

        col1, col2 = st.columns(2)
        
        with col1:
            st.subheader("📄 Export Document")
            doc_file = convert_html_to_doc_format(content)
            st.download_button(
                "📥 Télécharger la réponse (.doc)",
                doc_file,
                "Reponse_Sofia.doc",
                "application/msword"
            )
            
        with col2:
            st.subheader("📚 Sources PDF")
            if st.button("Préparer le ZIP"):
                soup = BeautifulSoup(content, 'html.parser')
                sources = soup.find_all('div', class_='source-card')
                zip_buffer = io.BytesIO()
                with zipfile.ZipFile(zip_buffer, "w") as zf:
                    for i, source in enumerate(sources):
                        link = source.find('h2', class_='card-title').find('a')
                        if link and link.get('href'):
                            try:
                                resp = requests.get(link['href'], timeout=10)
                                name = f"{i+1}_{link.get_text()[:50].strip()}.pdf".replace('/', '_')
                                zf.writestr(name, resp.content)
                            except:
                                pass
                st.download_button(
                    "📥 Télécharger le .zip",
                    zip_buffer.getvalue(),
                    "Sources.zip",
                    "application/zip"
                )

# --- ONGLET 3 : AIDE AU PROMPT "CONTRAINTES" ---
with tab3:
    st.header("3.Cadrage du problème et contraintes")
    st.info("En remplissant les champs ci dessous, vous précisez le champs de contrainte qui s'applique sur votre problématique. Ces informations seront mises en forme dans un document .docx. pour être ensuite ajoutées au document généré par SofIA.")
    c1 = st.text_area("Peut on réduire le périmètre du problème :")
    c2 = st.radio("Type de problème :", ["Compliqué", "Complexe", "Pernicieux (Wicked)[lien wikipedia](https://en.wikipedia.org/wiki/Wicked_problem)"])
    c3 = st.text_area("Quels sont les acteurs vraiment concernés par le problème ? [pain point](https://marketpedia.ca/lexique/pain-points/#:~:text=En%20marketing%2C%20l'expression%20anglaise,dans%20leur%20exp%C3%A9rience%20d'achat.) :")
    c4 = st.text_area("Avez vous identifié des partenaires obligatoires ou des compétiteurs :")
    c5 = st.text_area("Avez vous identifié des Bénéfices tangibles :")
    c6 = st.text_area("Avez vous identifié des Bénéfices intangibles :")
    c7 = st.text_area("Avez vous des Objectifs chiffrés et [opposables](https://www.icopilots.com/discours-il-opposable/) :")
    c8 = st.text_area("Budget prévisionnel :")
    c9 = st.text_area("Planning et jalons :")
    c10 = st.text_area("Y a t-il des restrictions en terme de Communication ? ou des conseils ? :")
    c11 = st.text_area("Avez vous identifié des Vecteurs marketing :")
    c12 = st.text_area("Informations complémentaires :")

    # On mémorise les contraintes pour l'onglet 4
    st.session_state.cadrage = {
        "Périmètre": c1, "Nature": c2, "Douleurs": c3, "Partenaires": c4,
        "Bénéfices T": c5, "Bénéfices I": c6, "Objectifs": c7,
        "Budget": c8, "Planning": c9, "Com": c10, "Marketing": c11, "Infos": c12
    }

    if st.button("Générer le document de cadrage (.docx)"):
        prompt_doc = Document()
        prompt_doc.add_heading("Cadrage du Problème & Prompt pour Eval", 0)

        p = prompt_doc.add_paragraph("Ce document est à fournir à l'Assistant Eval. Se connecter à ")
        add_hyperlink(
            p,
            "https://m365.cloud.microsoft/chat/?titleId=T_7b923e69-c9aa-4317-d331-4647b285be26",
            "Eval"
        )
       
        for key, value in st.session_state.cadrage.items():
            prompt_doc.add_heading(key, level=1)
            prompt_doc.add_paragraph(value if value else "Non précisé")
        
        prompt_buffer = io.BytesIO()
        prompt_doc.save(prompt_buffer)
        prompt_buffer.seek(0)
        
        st.download_button(
            label="📥 Télécharger votre document de cadrage",
            data=prompt_buffer,
            file_name="Cadrage_Projet_Sofia.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# --- ONGLET 4 : EVAL IA (FUSION) ---
with tab4:
    st.header("🤖 1er Agent EVAL")
    st.info(
        "L'Agent EVAL propose des modes d'intervention à partir d'une problématique décrite dans une pièce jointe; Il s'appuie sur une base documentaire des 100 programmes d'évaluation d'interventions de l'ADEME."
    )
    st.info(
        "Il faut fournir à cet Agent EVAL l'analyse SofIA (onglet 2) et vos contraintes de cadrage (onglet 3) "
        "dans un seul document optimisé pour l'IA."
    )

    gem_url = "https://gemini.google.com/gem/1y9w9p-YCpKER7F9enlRczopToIylE-NP?usp=sharing"

    if st.session_state.get("sofia_html") is None:
        st.warning("Veuillez d'abord importer le fichier chat_history.html dans l'onglet 2.")
    else:
        if st.button("Générer le document fusionné (.doc)"):
            merged_file = merge_html_and_constraints(
                st.session_state.sofia_html,
                st.session_state.cadrage
            )
            st.session_state.merged_file = merged_file
            st.session_state.fusion_generee = True

        if st.session_state.get("fusion_generee", False) and st.session_state.get("merged_file") is not None:
            st.download_button(
                label="📥 Télécharger le document fusionné",
                data=st.session_state.merged_file,
                file_name="Analyse_Sofia_et_Cadrage.doc",
                mime="application/msword"
            )

            st.markdown("---")
            st.success("✅ Document fusionné généré avec succès.")

            # Bouton lien vers GEM (ouvre dans le même onglet, l’utilisateur peut Ctrl+clic pour nouvel onglet)
            st.link_button("➡️ Aller vers l'Agent EVAL pour poursuivre l'analyse", gem_url)
            st.success("✅ Fournir le document généré à l'Agent EVAL en pièce jointe et copier/coller le Prompt suivant : *Pour résoudre le problème présenté dans le document joint, atteindre les objectifs fixés dans le contexte présenté, impliquer les acteurs identifiés, donne les modes d'intervention les plus adaptés. Il est possible de combiner plusieurs modes. Fournit les justifications précises pour choix de modes d'intervention et de combinaisons, et indique pourquoi les autres ne conviennent pas. Commence par résumer le problème à résoudre.*")
            st.success("📁 Une fois que l'Agent EVAL a répondu, en bas de la réponse, à coté de 👍 👎 🔄 📑, clique sur 3 petits points verticaux ⁝ puis *Exportez vers Docs*. En bas d'écran, cela génère un Google Doc que vous transformez via *Fichier / Télécharger / Microsoft Word*")
            st.warning("❤️ Il est possible et souhaitable d'évaluer le résultat fourni en allant ensuite à l'onglet 6.")

# --- ONGLET 5 : CHAOS IA ---
with tab5:
    st.header("🤖 2ème Agent CHAOS")
    st.info(
        "L'Agent Chaos vise à renforcer la robustesse d'un mode d'intervention en le stressant avec plusieurs évènements probables."
    )
    st.info(
        "Il faut fournir en pièce jointe à cet Agent CHAOS le document (.docx) généré par l'Agent EVAL (onglet 4). Si vous ne l'avez pas fait : en bas de la réponse de l'Agent EVAL, à coté de 👍 👎 🔄 📑, clique sur 3 petits points verticaux ⁝ puis *Exportez vers Docs*. En bas d'écran, cela génère un Google Doc que vous transformez via *Fichier / Télécharger / Microsoft Word*). "
    )

    gem2_url = "https://gemini.google.com/gem/1MdHeIqf8zOuSRWyO3RudoIHji3njSbtR?usp=sharing"

 # Bouton lien vers GEM (ouvre dans le même onglet, l’utilisateur peut Ctrl+clic pour nouvel onglet)
    st.link_button("➡️ Aller vers l'Agent CHAOS pour poursuivre l'analyse", gem2_url)
    st.success("✅ Fournir le document généré par l'Agent EVAL à l'Agent CHAOS en pièce jointe et copier/coller le Prompt suivant : *Evalue la robustesse du programme présenté dans le fichier ci-joint en sélectionnant des évènements probables à fort impacts et propose des améliorations sur l'organisation et les actions du programme, ainsi que les acteurs à impliquer*.")
    st.success("📁 Une fois que l'Agent CHAOS a répondu, en bas de la réponse, à coté de 👍 👎 🔄 📑, clique sur 3 petits points verticaux ⁝ puis *Exportez vers Docs*. En bas d'écran, cela génère un Google Doc que vous transformez via *Fichier / Télécharger / Microsoft Word*")
    st.warning("❤️ Il est possible et souhaitable d'évaluer le résultat fourni en allant ensuite à l'onglet 6.")

# --- ONGLET 6 : QUESTIONNAIRE ---
with tab6:
    st.header("6.Questionnaire d'évaluation des réponses des Agents")

    airtable_iframe = """
    <iframe 
        class="airtable-embed" 
        src="https://airtable.com/embed/apprjFOnbLySO8spa/shrncKacF4alPuLNZ" 
        frameborder="0" 
        onmousewheel="" 
        width="100%" 
        height="533" 
        style="background: transparent; border: 1px solid #ccc;">
    </iframe>
    """

# Intégration dans Streamlit
    components.html(airtable_iframe, height=550) # On met un height légèrement supérieur pour éviter les scrolls inutiles


    
